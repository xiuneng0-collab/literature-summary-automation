#!/usr/bin/env python
"""Convert the standard paper-summary Markdown file to DOCX with images embedded."""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


FONT_CN = "\u5b8b\u4f53"
FONT_HEAD = "\u5fae\u8f6f\u96c5\u9ed1"


def clean_inline(s: str) -> str:
    s = re.sub(r"`([^`]*)`", r"\1", s.strip())
    return s.replace("**", "")


def add_inline_runs(paragraph, s: str):
    pos = 0
    pattern = re.compile(r"(\*\*.*?\*\*|`.*?`)")
    for m in pattern.finditer(s):
        if m.start() > pos:
            paragraph.add_run(s[pos : m.start()])
        token = m.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(80, 80, 80)
        pos = m.end()
    if pos < len(s):
        paragraph.add_run(s[pos:])


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, value: str, bold: bool = False):
    cell.text = ""
    p = cell.paragraphs[0]
    add_inline_runs(p, value)
    for run in p.runs:
        run.font.name = FONT_CN
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        run.font.size = Pt(9.5)
        if bold:
            run.bold = True
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, md_table_lines):
    rows = []
    for line in md_table_lines:
        if re.match(r"^\|\s*-+", line):
            continue
        rows.append([p.strip() for p in line.strip().strip("|").split("|")])
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    rows = [r + [""] * (max_cols - len(r)) for r in rows]
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_text(cell, value, bold=(r_idx == 0 or c_idx == 0))
            if r_idx == 0:
                shade_cell(cell, "D9EAF7")
            elif c_idx == 0:
                shade_cell(cell, "F2F3F5")
    doc.add_paragraph()


def resolve_img_path(summary_path: Path, raw: str) -> Path:
    normalized = raw.replace("\\", "/")
    if "/outputs/" in normalized:
        # Prefer absolute path as-is when the host can resolve it.
        return Path(raw.replace("/", "\\"))
    p = Path(raw)
    if p.is_absolute():
        return p
    return summary_path.parent / raw


def make_docx(summary_path: Path, output_path: Path):
    text = summary_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    styles = doc.styles
    styles["Normal"].font.name = FONT_CN
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    styles["Normal"].font.size = Pt(10.5)
    for name, size in [("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)]:
        styles[name].font.name = FONT_HEAD
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEAD)
        styles[name].font.size = Pt(size)
        styles[name].font.bold = True

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table(doc, table_lines)
            continue
        img_match = re.match(r"^!\[(.*?)\]\((.*?)\)$", line.strip())
        if img_match:
            alt, img = img_match.groups()
            img_path = resolve_img_path(summary_path, img)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if img_path.exists():
                width = Inches(6.25)
                name = img_path.name.lower()
                if "table10" in name:
                    width = Inches(6.5)
                elif "table9" in name or "fig4" in name or "fig1" in name:
                    width = Inches(4.7)
                p.add_run().add_picture(str(img_path), width=width)
                cap = doc.add_paragraph(clean_inline(alt))
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cap.runs:
                    run.font.size = Pt(9)
                    run.italic = True
            else:
                add_inline_runs(p, "[图片未找到] " + str(img_path))
            i += 1
            continue
        if line.startswith("# "):
            p = doc.add_heading("", level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline_runs(p, line[2:].strip())
            for run in p.runs:
                run.font.name = FONT_HEAD
                run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEAD)
                run.font.size = Pt(20)
                run.bold = True
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(clean_inline(line[3:]), level=2)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(clean_inline(line[4:]), level=3)
            i += 1
            continue
        if line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.right_indent = Inches(0.25)
            add_inline_runs(p, line[2:].strip())
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(90, 90, 90)
            i += 1
            continue
        if re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, re.sub(r"^\d+\.\s+", "", line))
            i += 1
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, line[2:].strip())
            i += 1
            continue
        p = doc.add_paragraph()
        add_inline_runs(p, line)
        i += 1
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main():
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    ap = argparse.ArgumentParser()
    ap.add_argument("--summary", required=True, type=Path)
    ap.add_argument("--output", required=True, type=Path)
    args = ap.parse_args()
    make_docx(args.summary, args.output)
    print(args.output)


if __name__ == "__main__":
    main()
